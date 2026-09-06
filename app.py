import streamlit as st
import streamlit.components.v1 as components
import google.generativeai as genai
from PIL import Image
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import io
import urllib.parse
import time
import json
import base64

st.set_page_config(page_title="AdShield | Kurumsal Denetim", layout="wide", initial_sidebar_state="collapsed")

def eklenti_verilerini_getir():
    try:
        u_url = st.secrets.get("UPSTASH_REDIS_REST_URL")
        u_token = st.secrets.get("UPSTASH_REDIS_REST_TOKEN")
        if not u_url or not u_token: return []
            
        headers = {"Authorization": f"Bearer {u_token}"}
        res = requests.get(f"{u_url}/lrange/adshield_queue/0/-1", headers=headers, timeout=5)
        if res.status_code == 200:
            data = res.json()
            raw_list = data.get("result", [])
            if raw_list:
                requests.get(f"{u_url}/del/adshield_queue", headers=headers, timeout=3)
                islenmis_liste = []
                for item in raw_list:
                    parsed = json.loads(item)
                    if isinstance(parsed, str): parsed = json.loads(parsed)
                    islenmis_liste.append(parsed)
                return islenmis_liste
    except: pass
    return []

def create_single_docx(metin, gorsel=None, url="", is_dilekce=False):
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    if url and not is_dilekce:
        p_url = doc.add_paragraph()
        p_url.add_run("İhlal Kaynağı (URL): ").bold = True
        p_url.add_run(url)
        
    if gorsel and not is_dilekce:
        try:
            img_io = io.BytesIO()
            gorsel.save(img_io, format="PNG")
            img_io.seek(0)
            doc.add_picture(img_io, width=Inches(4.5))
        except: pass

    for line in metin.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith("--- DİLEKÇE"): continue

        if is_dilekce and line in ["T.C. TİCARET BAKANLIĞI", "REKLAM KURULU BAŞKANLIĞINA", "ANKARA"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line).bold = True
            continue

        if line.startswith('### '): doc.add_heading(line[4:].replace('**', ''), level=3)
        elif line.startswith('## '): doc.add_heading(line[3:].replace('**', ''), level=2)
        elif line.startswith('# '): doc.add_heading(line[2:].replace('**', ''), level=1)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            parts = line[2:].split('**')
            for i, part in enumerate(parts):
                if part: p.add_run(part).bold = (i % 2 != 0)
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if part: p.add_run(part).bold = (i % 2 != 0)
                
    docx_io = io.BytesIO()
    doc.save(docx_io)
    return docx_io.getvalue()

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@600;700&family=Plus+Jakarta+Sans:wght@400;500;600&display=swap');
    * { font-family: 'Plus Jakarta Sans', sans-serif; }
    .block-container { max-width: 1200px !important; padding-top: 1.2rem !important; margin: 0 auto !important; }
    .firm-header { background-color: #2C3848; padding: 18px 26px; border-radius: 6px; color: #ffffff; display: flex; justify-content: space-between; align-items: center; margin-bottom: 18px; }
    .firm-title { font-family: 'Cinzel', serif; font-size: 19px; font-weight: 700; letter-spacing: 1.5px; }
    .firm-subtitle { font-size: 11px; color: #DCE4EC; margin-top: 2px; }
    .firm-badge { background: rgba(255, 255, 255, 0.15); border: 1px solid rgba(255, 255, 255, 0.3); font-size: 11px; padding: 5px 12px; border-radius: 4px; }
    .section-heading { font-family: 'Cinzel', serif; font-size: 13.5px; font-weight: 700; color: #2C3848; border-bottom: 2px solid #E2E8F0; padding-bottom: 5px; margin-bottom: 12px; }
    div[data-testid="stRadio"] > div[role="radiogroup"] { display: flex; justify-content: center; gap: 14px; width: 100%; margin-bottom: 20px; }
    div[data-testid="stRadio"] > div[role="radiogroup"] > label { flex: 1; background: #FFFFFF; border: 1.5px solid #CBD5E1; border-radius: 6px; padding: 12px 18px; cursor: pointer; text-align: center; display: flex; align-items: center; justify-content: center; }
    div[data-testid="stRadio"] > div[role="radiogroup"] > label > div:first-child { display: none !important; }
    div[data-testid="stRadio"] > div[role="radiogroup"] > label:hover { border-color: #5D728B; background: #F8FAFC; }
    div[data-testid="stRadio"] > div[role="radiogroup"] > label:has(input:checked) { border-color: #5D728B !important; background-color: #F1F5F9 !important; font-weight: 600 !important; color: #1E293B !important; }
</style>
<div class="firm-header" lang="tr">
    <div><div class="firm-title">ADSHIELD PRO</div><div class="firm-subtitle">Hukuki Mütalaa & Rekabet Taarruz Sistemi</div></div>
    <div class="firm-badge">Tam Korumalı & Çift Çıktı Motoru Aktif</div>
</div>
""", unsafe_allow_html=True)

api_key = st.secrets.get("GEMINI_API_KEY", "")
SABIT_SERP_KEY = st.secrets.get("SERPAPI_API_KEY", "")

with st.sidebar:
    st.header("Sistem Ayarları")
    api_key_input = st.text_input("Gemini API Key:", value=api_key, type="password")
    serpapi_key_input = st.text_input("SerpApi Key:", value=SABIT_SERP_KEY, type="password")
    api_key = api_key_input if api_key_input else api_key
    SABIT_SERP_KEY = serpapi_key_input if serpapi_key_input else SABIT_SERP_KEY

def resize_for_api(image, max_size=1024):
    img = image.copy()
    img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
    return img

def get_master_prompt(sektor, mecra, metin="", url=""):
    ek_metin = f"\nKullanıcı Tarafından Belirtilen Ek Metin/İddia: {metin}\n" if metin else ""
    ek_url = f"İhlal Linki: {url}\n" if url else ""
    return f"""SEN TİCARET BAKANLIĞI REKLAM KURULU BAŞDENETÇİSİ VE UZMAN BİR HAKSIZ REKABET AVUKATISIN.
Sektör: {sektor} | Mecra: {mecra} 
{ek_url}{ek_metin}

Sana ekte bir reklam/tanıtım görseli (veya ekran görüntüsü) sunulmuştur. Bu görseldeki tüm metinleri, ince yazıları (dipnotları), ürün üzerindeki etiket iddialarını ve görselin yarattığı genel tüketici algısını kendin okuyup analiz edeceksin. YAPAY ZEKA AĞZI KULLANMADAN, resmi hukuk dilinde İKİ FARKLI METİN üret. Araya KESİNLİKLE "--- DİLEKÇE BAŞLANGICI ---" ayıracını koy.

[BÖLÜM 1: MÜTALAA (Görsel ve Rekabet Analizi)]
Görselden elde ettiğin bulguları ve (varsa) ek metni hukuki dayanaklarıyla (6502 sayılı Kanun m.61, Ticari Reklam Yön. m.7 vb.) açıkla. 
- Görseldeki "Klinik olarak kanıtlanmıştır", "1 Numara", "%100 etkili" gibi ispat külfeti gerektiren mutlak iddiaları acımadan işaretle.
- Rakipleri doğrudan veya dolaylı yoldan kötüleyen, onlara kıyasla haksız avantaj sağlayan üstünlük iddialarını (haksız rekabet) deşifre et.
- Sağlık beyanlarını ve "içermez" hilelerini tavizsiz incele.
- Görseldeki ana vaat ile okunması zor alt yazılar (istisnalar) arasında çelişki varsa bunu vurgula.

--- DİLEKÇE BAŞLANGICI ---
(BU KISIMDAN SONRA ASLA MADDELENDİRME İŞARETİ (*) KULLANMA. NUMARALI BAŞLIK (1., 2.) KULLAN.)

T.C. TİCARET BAKANLIĞI
REKLAM KURULU BAŞKANLIĞINA
ANKARA

ŞİKAYET EDEN 		: [Boş Bırak]
ADRES	 		: [Boş Bırak]
ŞİKAYET EDİLEN 	: [Görselden Firma Unvanını veya Markayı Çıkar]
ŞİKAYET KONUSU 	: Ekteki görselde ve satış sayfasında yer alan iddialar hakkında 6502 sayılı Tüketicinin Korunması Hakkında Kanun ve Ticari Reklam ve Haksız Ticari Uygulamalar Yönetmeliği uyarınca reklamların durdurulması ve ilgililer hakkında idari para cezası uygulanması talebidir.
AÇIKLAMALAR :
Şikayet edilen satıcı ve marka sahibi tarafından arz edilen ürünün görsel ve reklam materyalleri incelendiğinde; yürürlükteki mevzuat hükümleri ve Reklam Kurulu’nun yerleşik içtihatları çerçevesinde açıkça hukuka aykırılık teşkil ettiği tespit edilmiştir.

1. [GÖRSELDEN TESPİT EDİLEN ANA İHLAL BAŞLIĞI - BÜYÜK HARFLE]
[Görseldeki tam ifadeleri alıntılayarak ve hukuki dayanaklarla açıklama]

2. [DİĞER İHLAL/HAKSIZ REKABET BAŞLIĞI - BÜYÜK HARFLE]
[Hukuki dayanaklarla açıklama]

SONUÇ VE İSTEM	: Yukarıdaki açıklamalar çerçevesinde reklam ve bilgilendirmelerin incelenerek yayının durdurulmasına ve sorumlu şirketin idari para cezası ile cezalandırılmasına karar verilmesini talep ederiz."""

def ai_istek_at(prompt, gorsel, is_stream=False):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-3.0-pro")
    optimize_gorsel = resize_for_api(gorsel)
    
    if is_stream:
        return model.generate_content([prompt, optimize_gorsel], stream=False, generation_config={"temperature": 0.2})
    else:
        return model.generate_content([prompt, optimize_gorsel], generation_config={"temperature": 0.2}).text

def generate_multi_role_synthesis_stream(gorsel, sektor, mecra, metin="", url=""):
    prompt = get_master_prompt(sektor, mecra, metin, url)
    try:
        response = ai_istek_at(prompt, gorsel, is_stream=True)
        if response and response.text:
            words = response.text.split(' ')
            for i in range(0, len(words), 6):
                yield ' '.join(words[i:i+6]) + ' '
                time.sleep(0.02)
            return
    except Exception as e:
        yield f"\nSistem Hatası: {str(e)}"
        return

def tekil_sorgu_at(kategori, sorgu, api_key_val):
    url = f"https://serpapi.com/search.json?q={urllib.parse.quote(sorgu)}&api_key={api_key_val}&engine=google&gl=tr&hl=tr&num=20"
    try:
        response = requests.get(url, timeout=20)
        data = response.json()
        if "error" in data: return kategori, [{"baslik": "Hata/Engellendi", "url": "#", "snippet": data['error']}]
        link_havuzu = []
        for result in data.get("organic_results", []):
            link = result.get("link", "")
            if not link.startswith("http"): continue
            link_havuzu.append({"baslik": result.get("title", "İsimsiz"), "url": link, "snippet": result.get("snippet", "")})
        if not link_havuzu: return kategori, [{"baslik": "Uygun sonuç bulunamadı.", "url": "#", "snippet": ""}]
        return kategori, link_havuzu
    except Exception as e: 
        return kategori, [{"baslik": "⚠️ HATA", "url": "#", "snippet": str(e)}]

def gelismis_coklu_hedef_taramasi(urun_adi, api_key_val):
    if not api_key_val or not urun_adi.strip(): return {}
    t = urun_adi.strip()
    queries = {
        "📸 Instagram": f'site:instagram.com "{t}"',
        "🛒 E-Ticaret": f'(site:trendyol.com OR site:hepsiburada.com) "{t}"'
    }
    kategorize_sonuclar = {}
    for kat, q in queries.items():
        _, sonuclar = tekil_sorgu_at(kat, q, api_key_val)
        gorulenler = set()
        tekil_list = []
        for item in sonuclar:
            if item["url"] not in gorulenler:
                gorulenler.add(item["url"])
                tekil_list.append(item)
        kategorize_sonuclar[kat] = tekil_list
    return kategorize_sonuclar

for k in ["rapor_sonucu", "eklenti_img", "eklenti_url", "vaka_havuzu", "radar_link_sonuclari"]:
    if k not in st.session_state: st.session_state[k] = None if k != "vaka_havuzu" else []

mod_secimi = st.radio("Denetim Modu", [
    "Kurumsal Kampanya Taslağı Uyum Denetimi",
    "Piyasa ve Rakip Reklam İncelemesi",
    "360° Çoklu Satıcı ve Pazar Radarı"
], horizontal=True, label_visibility="collapsed")

is_radar = "360°" in mod_secimi

sol_kolon, sag_kolon = st.columns([1, 1.25], gap="large")

if not is_radar:
    with sol_kolon:
        with st.container(border=True):
            st.markdown('<div class="section-heading" lang="tr">1. Veri Yükleme</div>', unsafe_allow_html=True)
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📥 Son Görüntüyü Al"):
                    gelen = eklenti_verilerini_getir()
                    if gelen:
                        son = gelen[-1]
                        st.session_state.eklenti_url = son.get('url', '')
                        b64 = son.get('image', '').split(',')[-1].strip() + '=='
                        st.session_state.eklenti_img = Image.open(io.BytesIO(base64.b64decode(b64))).convert("RGB")
                        st.rerun()

            if st.session_state.eklenti_img:
                st.image(st.session_state.eklenti_img, use_container_width=True)
                sektor = st.selectbox("Faaliyet Sektörü", ["Kozmetik & Kişisel Bakım", "Takviye Edici Gıda & Sağlık", "E-Ticaret", "Diğer"])
                mecra = st.selectbox("Yayınlanacak Mecra", ["İnternet / Sosyal Medya", "Satış Noktası", "Televizyon", "Açık Hava"])
                reklam_url = st.text_input("Web Sayfası / Ürün Linki", value=st.session_state.get("eklenti_url", ""))
                reklam_metni = st.text_area("Reklam Metni / Ticari İddialar", height=90)
                
                if st.button("Hukuki Analizi & Dilekçeyi Başlat", type="primary"):
                    st.session_state.rapor_sonucu = None
            
    with sag_kolon:
        with st.container(border=True):
            st.markdown('<div class="section-heading" lang="tr">2. Çıktılar</div>', unsafe_allow_html=True)
            
            if st.session_state.eklenti_img and st.session_state.rapor_sonucu is None and reklam_url is not None:
                rapor_alani = st.empty()
                tam_rapor = ""
                for parca in generate_multi_role_synthesis_stream(st.session_state.eklenti_img, sektor, mecra, reklam_metni, reklam_url):
                    tam_rapor += parca
                    rapor_alani.markdown(tam_rapor + "▌")
                rapor_alani.markdown(tam_rapor)
                st.session_state.rapor_sonucu = tam_rapor
                
            elif st.session_state.rapor_sonucu:
                st.markdown(st.session_state.rapor_sonucu)
                
            if st.session_state.rapor_sonucu:
                rapor_kismi = st.session_state.rapor_sonucu
                dilekce_kismi = ""
                if "--- DİLEKÇE BAŞLANGICI ---" in st.session_state.rapor_sonucu:
                    parcalar = st.session_state.rapor_sonucu.split("--- DİLEKÇE BAŞLANGICI ---")
                    rapor_kismi = parcalar[0]
                    if len(parcalar) > 1:
                        dilekce_kismi = parcalar[1]
                
                st.divider()
                col_indir_1, col_indir_2 = st.columns(2)
                with col_indir_1:
                    word_rapor = create_single_docx(rapor_kismi, gorsel=st.session_state.eklenti_img, url=reklam_url, is_dilekce=False)
                    st.download_button("📊 Mütalaa Raporunu İndir", word_rapor, "adshield_mutalaa.docx", use_container_width=True)
                with col_indir_2:
                    if dilekce_kismi:
                        word_dilekce = create_single_docx(dilekce_kismi, is_dilekce=True)
                        st.download_button("⚖️ Şikayet Dilekçesini İndir", word_dilekce, "adshield_dilekce.docx", use_container_width=True, type="primary")

else:
    with sol_kolon:
        with st.container(border=True):
            st.markdown('<div class="section-heading" lang="tr">Pazar Radarı</div>', unsafe_allow_html=True)
            radar_urun = st.text_input("Marka / İhlal Kelimesi", value="incia bebek yağı")
            if st.button("Hedefleri Bul", type="primary"):
                with st.spinner("Tarama sürüyor..."):
                    st.session_state.radar_link_sonuclari = gelismis_coklu_hedef_taramasi(radar_urun, SABIT_SERP_KEY)
    with sag_kolon:
        with st.container(border=True):
            st.markdown('<div class="section-heading" lang="tr">Bulgular</div>', unsafe_allow_html=True)
            if st.session_state.get("radar_link_sonuclari"):
                sekmeler = st.tabs(list(st.session_state.radar_link_sonuclari.keys()))
                for i, (kat, links) in enumerate(st.session_state.radar_link_sonuclari.items()):
                    with sekmeler[i]:
                        for idx, item in enumerate(links, 1): st.markdown(f"**{idx}. [{item['baslik']}]({item['url']})**")
